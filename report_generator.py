"""
DOCX报告生成器 - 按党政公文格式 (GB/T 9704-2012) 输出

字体规范:
- 标题: 方正小标宋简体, 二号 (22pt), 居中
- 一级标题: 黑体, 三号 (16pt)
- 二级标题: 楷体_GB2312, 三号 (16pt)
- 正文: 仿宋_GB2312, 三号 (16pt)
"""
import io
import json
import logging
from collections import Counter
from datetime import datetime, timedelta, timezone, date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from sqlalchemy import func, desc
from sqlalchemy.orm import Session

from models import Article, TopicCluster, TrendSnapshot
from analyzer import extract_keywords

logger = logging.getLogger(__name__)

# 字体常量 (公文标准)
FONT_TITLE = "方正小标宋简体"
FONT_H1 = "黑体"
FONT_H2 = "楷体_GB2312"
FONT_BODY = "仿宋_GB2312"

# 字号 (pt)
SIZE_TITLE = 22   # 二号
SIZE_H1 = 16      # 三号
SIZE_H2 = 16      # 三号
SIZE_BODY = 16    # 三号

SOURCE_NAMES = {
    "gdelt": "GDELT 全球事件数据库",
    "google_news": "Google News",
    "mediacloud": "Media Cloud",
}

SENTIMENT_NAMES = {
    "positive": "正面",
    "negative": "负面",
    "neutral": "中性",
}


def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False):
    """为 run 设置字体 (中英文都设置, 并正确处理 eastAsia 字体)."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # 关键: 设置 rFonts 的 eastAsia 属性, 确保中文使用该字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _add_title(doc, text: str):
    """标题 - 方正小标宋, 二号, 居中."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(text)
    _set_run_font(run, FONT_TITLE, SIZE_TITLE, bold=False)
    return p


def _add_h1(doc, text: str):
    """一级标题 - 黑体, 三号."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(SIZE_H1 * 2)  # 首行缩进2字符
    run = p.add_run(text)
    _set_run_font(run, FONT_H1, SIZE_H1, bold=False)
    return p


def _add_h2(doc, text: str):
    """二级标题 - 楷体_GB2312, 三号."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Pt(SIZE_H2 * 2)
    run = p.add_run(text)
    _set_run_font(run, FONT_H2, SIZE_H2, bold=False)
    return p


def _add_body(doc, text: str, indent: bool = True):
    """正文 - 仿宋_GB2312, 三号, 首行缩进2字符."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(SIZE_BODY * 2)
    run = p.add_run(text)
    _set_run_font(run, FONT_BODY, SIZE_BODY, bold=False)
    return p


def _set_default_style(doc):
    """设置文档默认样式 (仿宋_GB2312 三号)."""
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(SIZE_BODY)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:ascii"), FONT_BODY)
    rFonts.set(qn("w:eastAsia"), FONT_BODY)
    rFonts.set(qn("w:hAnsi"), FONT_BODY)


def _set_page_margins(doc):
    """A4 页边距: 上 3.7cm, 下 3.5cm, 左 2.8cm, 右 2.6cm (公文标准)."""
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


def _fetch_stats(session: Session, days: int, keywords: list[str] | None = None):
    """收集统计数据."""
    cutoff = datetime.now(timezone.utc) - timedelta(days=days)
    q = session.query(Article).filter(Article.published_at >= cutoff)

    if keywords:
        from sqlalchemy import or_
        conds = []
        for kw in keywords:
            conds.append(Article.title.ilike(f"%{kw}%"))
            conds.append(Article.content_snippet.ilike(f"%{kw}%"))
            conds.append(Article.search_keywords.ilike(f"%{kw}%"))
        q = q.filter(or_(*conds))

    total = q.count()

    # 按来源
    source_rows = (
        q.with_entities(Article.source, func.count(Article.id))
        .group_by(Article.source).all()
    )
    by_source = dict(source_rows)

    # 情感分布
    sent_rows = (
        q.with_entities(Article.sentiment_label, func.count(Article.id))
        .filter(Article.sentiment_label.isnot(None))
        .group_by(Article.sentiment_label).all()
    )
    sentiment = {"positive": 0, "neutral": 0, "negative": 0}
    for label, cnt in sent_rows:
        if label in sentiment:
            sentiment[label] = cnt

    avg_sent = q.with_entities(func.avg(Article.sentiment_score)).scalar() or 0

    # 按国家
    country_rows = (
        q.with_entities(Article.country, func.count(Article.id))
        .filter(Article.country.isnot(None), Article.country != "")
        .group_by(Article.country)
        .order_by(desc(func.count(Article.id)))
        .limit(10).all()
    )
    by_country = list(country_rows)

    # 每日趋势
    daily_rows = (
        q.with_entities(func.date(Article.published_at).label("day"),
                        func.count(Article.id).label("cnt"))
        .group_by(func.date(Article.published_at))
        .order_by(func.date(Article.published_at)).all()
    )
    daily = [(str(r.day), r.cnt) for r in daily_rows if r.day]

    # 热点文章 (近期)
    top_articles = (
        q.order_by(desc(Article.published_at)).limit(10).all()
    )

    return {
        "total": total,
        "by_source": by_source,
        "sentiment": sentiment,
        "avg_sent": avg_sent,
        "by_country": by_country,
        "daily": daily,
        "top_articles": top_articles,
    }


def generate_report(session: Session, days: int = 30, keywords: list[str] | None = None) -> bytes:
    """
    生成 DOCX 报告, 返回二进制内容.

    :param session: SQLAlchemy session
    :param days: 统计天数
    :param keywords: 可选, 若提供则为关键词专题报告
    :return: .docx 文件二进制内容
    """
    doc = Document()
    _set_default_style(doc)
    _set_page_margins(doc)

    # === 标题 ===
    if keywords:
        title = f"关于「{' · '.join(keywords)}」舆情监测专题报告"
    else:
        title = "甪端全球新闻大数据监测分析报告"
    _add_title(doc, title)

    # === 报告信息 (署名+日期) ===
    now_str = datetime.now().strftime("%Y年%m月%d日")
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(18)
    info_run = info_p.add_run(f"甪端全球新闻大数据监测平台  {now_str}")
    _set_run_font(info_run, FONT_H2, 14, bold=False)

    # 分隔横线段落
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_p.add_run("—" * 30)
    _set_run_font(line_run, FONT_BODY, 12, bold=False)

    # 收集数据
    stats = _fetch_stats(session, days, keywords)

    # === 一、概述 ===
    _add_h1(doc, "一、监测概述")
    if keywords:
        overview = (
            f"本报告针对关键词「{'、'.join(keywords)}」, 基于甪端全球新闻大数据监测平台整合的"
            f"GDELT 全球事件数据库、Google News、Media Cloud 三大权威数据源, "
            f"对过去 {days} 天内的全球媒体报道进行综合分析。"
            f"监测期间共采集相关新闻 {stats['total']} 篇, 覆盖 {len(stats['by_country'])} 个国家和地区。"
        )
    else:
        overview = (
            f"本报告基于甪端全球新闻大数据监测平台, 整合 GDELT 全球事件数据库、"
            f"Google News 及 Media Cloud 三大权威数据源, "
            f"对过去 {days} 天内全球主流媒体的新闻报道进行多维度数据分析。"
            f"监测期间共采集新闻 {stats['total']} 篇, 覆盖 {len(stats['by_country'])} 个国家和地区。"
        )
    _add_body(doc, overview)

    # === 二、数据来源分布 ===
    _add_h1(doc, "二、数据来源分布")
    _add_body(doc, "本次监测涵盖以下数据源, 各数据源采集的新闻篇数如下:")
    if stats["by_source"]:
        total = max(stats["total"], 1)
        cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
        for idx, (src, cnt) in enumerate(sorted(stats["by_source"].items(), key=lambda x: -x[1])):
            name = SOURCE_NAMES.get(src, src)
            pct = cnt / total * 100
            num = cn_nums[idx] if idx < len(cn_nums) else str(idx + 1)
            _add_h2(doc, f"({num}){name}")
            _add_body(doc, f"共采集新闻 {cnt} 篇, 占总量 {pct:.1f}%。")
    else:
        _add_body(doc, "暂无数据。")

    # === 三、情感倾向分析 ===
    _add_h1(doc, "三、舆情情感倾向分析")
    sent = stats["sentiment"]
    sent_total = max(sum(sent.values()), 1)
    avg_sent = stats["avg_sent"]

    if avg_sent > 0.1:
        overall = "整体舆情呈正面倾向"
    elif avg_sent < -0.1:
        overall = "整体舆情呈负面倾向"
    else:
        overall = "整体舆情相对中性平稳"

    _add_body(doc,
        f"经情感分析模型处理, {overall}, "
        f"平均情感得分为 {avg_sent:.3f} (取值范围 -1 到 +1, 正值代表正面)。"
    )
    _add_body(doc,
        f"其中正面报道 {sent['positive']} 篇, 占比 {sent['positive']/sent_total*100:.1f}%; "
        f"中性报道 {sent['neutral']} 篇, 占比 {sent['neutral']/sent_total*100:.1f}%; "
        f"负面报道 {sent['negative']} 篇, 占比 {sent['negative']/sent_total*100:.1f}%。"
    )

    # === 四、地域分布 ===
    _add_h1(doc, "四、地域分布")
    if stats["by_country"]:
        _add_body(doc, f"报道覆盖国家和地区 Top 10 如下:")
        for i, (country, cnt) in enumerate(stats["by_country"], 1):
            _add_body(doc, f"{i}. {country}: {cnt} 篇")
    else:
        _add_body(doc, "暂无地域分布数据。")

    # === 五、热点话题 ===
    _add_h1(doc, "五、热点话题聚类")
    clusters = session.query(TopicCluster).order_by(desc(TopicCluster.article_count)).limit(8).all()
    if clusters:
        _add_body(doc, "基于 TF-IDF + KMeans 算法对近期新闻进行话题聚类, 主要热点话题如下:")
        for i, c in enumerate(clusters, 1):
            kws = json.loads(c.keywords) if c.keywords else []
            kw_str = "、".join(kws[:6]) if kws else "(无关键词)"
            _add_h2(doc, f"(话题 {i}){c.label}")
            _add_body(doc, f"相关报道 {c.article_count} 篇, 核心关键词: {kw_str}。")
    else:
        _add_body(doc, "暂无话题聚类数据, 请稍候系统自动分析。")

    # === 六、关键词词云分析 ===
    _add_h1(doc, "六、高频关键词分析")
    try:
        top_kws = extract_keywords(session, days=days, top_n=30)
    except Exception as e:
        logger.error(f"Keyword extraction error: {e}")
        top_kws = []
    if top_kws:
        _add_body(doc, "从新闻标题中提取的 Top 30 高频关键词如下 (按词频降序):")
        rows = []
        for i, kw in enumerate(top_kws, 1):
            rows.append(f"{kw['name']} ({kw['value']})")
        # 每行 5 个
        for i in range(0, len(rows), 5):
            _add_body(doc, "、".join(rows[i:i+5]) + ("。" if i+5 >= len(rows) else ";"))
    else:
        _add_body(doc, "暂无关键词数据。")

    # === 七、每日报道趋势 ===
    _add_h1(doc, "七、每日报道量趋势")
    if stats["daily"]:
        daily = stats["daily"]
        total_daily = sum(c for _, c in daily)
        avg_daily = total_daily / len(daily) if daily else 0
        peak = max(daily, key=lambda x: x[1]) if daily else ("", 0)
        _add_body(doc,
            f"监测期间日均报道量约 {avg_daily:.0f} 篇, "
            f"报道峰值出现在 {peak[0]}, 当日共 {peak[1]} 篇。"
        )
        _add_body(doc, "近 14 天每日报道量明细:")
        for day_str, cnt in daily[-14:]:
            _add_body(doc, f"  {day_str}: {cnt} 篇", indent=False)
    else:
        _add_body(doc, "暂无每日趋势数据。")

    # === 八、近期重点报道 ===
    _add_h1(doc, "八、近期重点报道")
    if stats["top_articles"]:
        _add_body(doc, "按发布时间筛选近期代表性报道如下:")
        for i, a in enumerate(stats["top_articles"][:10], 1):
            pub_str = a.published_at.strftime("%Y-%m-%d %H:%M") if a.published_at else "未知时间"
            sent_label = SENTIMENT_NAMES.get(a.sentiment_label or "", "未知")
            src_name = SOURCE_NAMES.get(a.source, a.source)
            _add_h2(doc, f"{i}. {a.title[:80]}")
            _add_body(doc,
                f"来源: {a.source_name or src_name}  "
                f"发布时间: {pub_str}  "
                f"情感: {sent_label}"
            )
            if a.content_snippet:
                snippet = a.content_snippet[:200].strip()
                _add_body(doc, f"摘要: {snippet}...")
    else:
        _add_body(doc, "暂无重点报道。")

    # === 九、结论与建议 ===
    _add_h1(doc, "九、结论")
    if keywords:
        conclusion = (
            f"综上所述, 关于「{'、'.join(keywords)}」的舆情监测结果显示: "
            f"过去 {days} 天内共产生相关报道 {stats['total']} 篇, "
            f"{overall}。建议持续关注相关话题动态, 及时研判舆情走向。"
        )
    else:
        conclusion = (
            f"综上所述, 过去 {days} 天全球新闻舆情总体{overall}。"
            f"建议结合地域分布、话题聚类、关键词及情感倾向等多维度数据, "
            f"持续跟踪热点议题, 加强舆情研判与应对。"
        )
    _add_body(doc, conclusion)

    # 落款
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(24)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run(f"甪端全球新闻大数据监测平台")
    _set_run_font(footer_run, FONT_H2, SIZE_BODY, bold=False)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_p.add_run(now_str)
    _set_run_font(date_run, FONT_H2, SIZE_BODY, bold=False)

    # 输出到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
